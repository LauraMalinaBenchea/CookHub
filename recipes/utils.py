import logging
import re
from decimal import Decimal

import docx
import fitz
import nltk
import spacy
from django.db import transaction

from recipes.constants import recipe_ingredient_keywords, recipe_step_keywords
from recipes.models import Ingredient, Recipe, RecipeIngredient, Step, Unit

nlp = spacy.load("en_core_web_sm")


nltk.download("punkt_tab")
logger = logging.getLogger(__name__)


def extract_text_from_file(file_obj: str, file_type: str) -> str:
    logger.info(f"Attempting to extract text from {file_type} file.")
    try:
        if file_type == ".pdf":
            return extract_text_from_pdf(file_obj)
        elif file_type == ".docx":
            return extract_text_from_docx(file_obj)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as exc:
        logger.error(f"Error extracting text from file: {exc}")


def extract_text_from_pdf(file_obj: str) -> str:
    text = ""
    file_bytes = file_obj.read()
    with fitz.open(stream=file_bytes) as doc:

        for page in doc:
            text += page.get_text()
        logger.info(f"Extracted text from PDF with {len(doc)} pages.")

    file_obj.seek(0)
    return text


def extract_text_from_docx(file_obj: str) -> str:
    doc = docx.Document(file_obj)
    text = "\n".join([para.text for para in doc.paragraphs])
    logger.info(
        f"Extracted text from DOCX with {len(doc.paragraphs)} paragraphs."
    )
    file_obj.seek(0)
    return text


def parse_ingredient_lines(ingredient_lines):
    # Load all units from DB
    all_units = list(Unit.objects.all())
    units = {u.abbreviation.lower(): u for u in all_units}
    units.update({u.name.lower(): u for u in all_units})

    # Fallback "piece" unit
    fallback_unit = Unit.objects.filter(abbreviation="pcs").first()

    ingredients_data = []

    quantity_pattern = re.compile(r"^\d+([\/\.\d]*)?")
    number_word_pattern = re.compile(
        r"\b(one|two|three|four|five|six|seven|eight|nine|ten)\b", re.I
    )

    for line in ingredient_lines:
        original_line = line
        line = line.strip().lower()

        # --- Extract numeric quantity ---
        match = quantity_pattern.match(line)
        if match:
            q_str = match.group(0)
            try:
                quantity = float(eval(q_str))  # supports fractions like 1/2
            except Exception:
                print(f"⚠️ Failed to parse quantity: {q_str}, fallback to 1.0")
                quantity = 1.0
            line = line[len(q_str) :].strip()
        else:
            # fallback for word-based quantities (e.g. "two eggs")
            word_match = number_word_pattern.search(line)
            if word_match:
                word_to_num = {
                    "one": 1,
                    "two": 2,
                    "three": 3,
                    "four": 4,
                    "five": 5,
                    "six": 6,
                    "seven": 7,
                    "eight": 8,
                    "nine": 9,
                    "ten": 10,
                }
                quantity = word_to_num[word_match.group(1).lower()]
                line = line.replace(word_match.group(0), "").strip()
            else:
                print("⚠️ Failed to parse quantity: None, fallback to 1.0")
                quantity = 1.0

        # --- Try to detect the unit ---
        unit = None
        for key, unit_obj in units.items():
            if re.search(rf"\b{re.escape(key)}\b", line):
                unit = unit_obj
                line = re.sub(rf"\b{re.escape(key)}\b", "", line).strip()
                break

        # --- Fallback to "pcs" if no unit found ---
        if not unit:
            unit = fallback_unit
            print(f"ℹ️ Fallback to unit 'pcs' for line: '{original_line}'")

        # --- Clean up ingredient name ---
        name = line.strip(" ,.-")

        ingredients_data.append(
            {
                "name": name,
                "quantity": Decimal(str(quantity)),
                "unit": unit,
                "raw": original_line,
            }
        )

    return ingredients_data


def parse_recipe_from_text(text: str, user, privacy="private") -> Recipe:
    """
    Parse raw text from a recipe file into a Recipe object and related models.
    """
    if not user:
        raise ValueError("User must be provided to create recipe.")

    # TODO: allow selecting privacy when uploading recipe
    # Normalize text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    # joined = "\n".join(lines).lower()

    title = lines[0] if lines else "Untitled Recipe"

    ingredients_idx = None
    steps_idx = None

    # TODO optimize understanding recipe data, to prevent
    #  fetching indexes based on recipe description
    ingredients_pattern = re.compile(
        r"|".join([re.escape(k) for k in recipe_ingredient_keywords]), re.I
    )
    steps_pattern = re.compile(
        r"|".join([re.escape(k) for k in recipe_step_keywords]), re.I
    )

    for i, line in enumerate(lines):
        if ingredients_pattern.search(line) and not ingredients_idx:
            ingredients_idx = i
            print("ingredients index ", ingredients_idx)
        if steps_pattern.search(line) and not steps_idx:
            steps_idx = i

            print("Steps index", steps_idx)

    if ingredients_idx is None:
        raise ValueError("No 'Ingredients' section found in text.")

    description_lines = []
    if ingredients_idx > 1:
        description_lines = lines[1:ingredients_idx]
    description = " ".join(description_lines)

    if steps_idx:
        ingredient_lines = lines[ingredients_idx + 1 : steps_idx]
        step_lines = lines[steps_idx + 1 :]
    else:
        ingredient_lines = lines[ingredients_idx + 1 :]
        step_lines = []

    ingredients_data = parse_ingredient_lines(ingredient_lines)
    print("Ingredients data", ingredients_data)

    steps = []
    for i, line in enumerate(step_lines):
        if re.match(r"^\d+[\).]", line):
            steps.append(line)
        elif steps:
            steps[-1] += " " + line
        else:
            steps.append(line)

    with transaction.atomic():
        recipe = Recipe.objects.create(
            title=title.strip(),
            description=description.strip(),
            privacy=privacy,
            servings=1,
            user=user,
        )

        for ing in ingredients_data:
            ingredient, _ = Ingredient.objects.get_or_create(
                name=ing["name"].lower()
            )
            RecipeIngredient.objects.create(
                recipe=recipe,
                ingredient=ingredient,
                quantity=ing["quantity"],
                unit=ing["unit"],
            )

        for idx, step_text in enumerate(steps, start=1):
            Step.objects.create(recipe=recipe, order=idx, text=step_text)

    return recipe


def convert_unit(quantity, from_unit, to_system):
    """
    Convert a quantity from one unit to another system (metric/imperial).
    Returns: (converted_quantity, target_unit_abbreviation)
    """
    # Exclude category "count"
    if from_unit.category not in ["weight", "volume"]:
        return quantity, from_unit.abbreviation

    # Get all units in the same category
    category_units = Unit.objects.filter(category=from_unit.category)

    # Find the target system base
    to_units = category_units.filter(system=to_system)
    if not to_units.exists():
        return quantity, from_unit.abbreviation

    # Try to find a "base" target unit, or fallback to the first
    to_unit = to_units.filter(is_base_unit=True).first() or to_units.first()
    base_quantity = quantity * from_unit.base_conversion_factor
    converted_quantity = base_quantity / to_unit.base_conversion_factor

    # Round neatly: 0 decimals if whole number, else 1 decimal
    converted_quantity = (
        round(converted_quantity, 1)
        if converted_quantity % 1
        else int(converted_quantity)
    )

    return converted_quantity, to_unit.abbreviation
