import docx
import fitz

from recipes.utils.recipe_processing import logger

"""
File used to keep logic related to extracting text from files.
"""


def extract_text_from_file(file_obj: str, file_type: str) -> str:
    # TODO: allow multiple recipes to come from a signle file?
    logger.info(f"Attempting to extract text from {file_type} file.")
    try:
        if file_type == ".pdf":
            return _extract_text_from_pdf(file_obj)
        elif file_type == ".docx":
            return _extract_text_from_docx(file_obj)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as exc:
        logger.error(f"Error extracting text from file: {exc}")


def _extract_text_from_docx(file_obj: str) -> str:
    doc = docx.Document(file_obj)
    text = "\n".join([para.text for para in doc.paragraphs])
    logger.info(
        f"Extracted text from DOCX with {len(doc.paragraphs)} paragraphs."
    )
    file_obj.seek(0)
    return text


def _extract_text_from_pdf(file_obj: str) -> str:
    text = ""
    file_bytes = file_obj.read()
    with fitz.open(stream=file_bytes) as doc:
        for page in doc:
            text += page.get_text()
        logger.info(f"Extracted text from PDF with {len(doc)} pages.")

    file_obj.seek(0)
    return text
